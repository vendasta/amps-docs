"""
Convert Word .docx to Markdown with images copied to Docusaurus static/img.
Preserves run order so inline images stay next to their text.
"""
from __future__ import annotations

import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.document import Document as DocumentClass
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def load_rels(docx_path: Path) -> dict[str, str]:
    with zipfile.ZipFile(docx_path) as z:
        rels_data = z.read("word/_rels/document.xml.rels")
    root = ET.fromstring(rels_data)
    rels: dict[str, str] = {}
    for rel in root:
        rid = rel.get("Id") or rel.get(
            "{http://schemas.openxmlformats.org/package/2006/relationships}Id"
        )
        target = rel.get("Target")
        if not rid or not target:
            continue
        if target.startswith("../"):
            target = "word/" + target[3:]
        elif not target.startswith("word/"):
            target = "word/" + target
        rels[rid] = target.replace("\\", "/")
    return rels


def iter_block_items(parent):
    if isinstance(parent, DocumentClass):
        parent_elm = parent.element.body
    else:
        raise ValueError("parent must be Document")
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def paragraph_run_chunks(p: Paragraph):
    """Yield ('image', rId) or ('text', str) in document order."""
    BLIP = "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
    for run in p.runs:
        for blip in run._element.findall(f".//{BLIP}"):
            embed = blip.get(qn("r:embed"))
            if embed:
                yield ("image", embed)
        if run.text:
            yield ("text", run.text)


def paragraph_plain_text(p: Paragraph) -> str:
    return "".join(run.text for run in p.runs)


def sanitize_line(s: str) -> str:
    return s.replace("\r", "").strip()


def is_junk_numbered_bullet(line: str) -> bool:
    t = line.strip()
    if not t:
        return False
    if re.match(r"^[\d]+\.\s*$", t):
        return True
    if re.match(r"^[\d]+\.\s*[\d]+\.\s*$", t):
        return True
    if re.match(r"^[ivxlcdm]+\)\s*$", t, re.I):
        return True
    if re.match(r"^[a-z]\)\s*$", t, re.I):
        return True
    return False


def docx_to_markdown(
    docx_path: Path,
    out_md: Path,
    static_img_subdir: str,
    doc_key: str,
    repo_root: Path,
    title_override: str | None = None,
) -> None:
    rels = load_rels(docx_path)
    dest_dir = repo_root / "docusaurus/static/img" / static_img_subdir
    dest_dir.mkdir(parents=True, exist_ok=True)

    image_counter = [0]

    def copy_image(rid: str) -> str | None:
        target = rels.get(rid)
        if not target:
            return None
        with zipfile.ZipFile(docx_path) as z:
            data = z.read(target)
        ext = Path(target).suffix.lower() or ".png"
        image_counter[0] += 1
        fname = f"{doc_key}-{image_counter[0]:02d}{ext}"
        out_path = dest_dir / fname
        out_path.write_bytes(data)
        return f"/img/{static_img_subdir}/{fname}".replace("\\", "/")

    doc = Document(docx_path)
    lines: list[str] = []

    title_guess = None
    for p in doc.paragraphs[:8]:
        t = paragraph_plain_text(p).strip()
        if t and title_guess is None and len(t) < 200:
            title_guess = t
            break

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            p = block
            style = p.style.name if p.style else ""
            full_text = sanitize_line(paragraph_plain_text(p))

            heading = False
            level = 2
            if style.startswith("Heading"):
                heading = True
                try:
                    level = int(style.replace("Heading", "").strip() or "1")
                except ValueError:
                    level = 1
                level = max(1, min(6, level))
            elif style == "Title":
                heading = True
                level = 1

            if heading and full_text:
                if is_junk_numbered_bullet(full_text):
                    continue
                lines.append("")
                lines.append(f"{'#' * level} {full_text}")
                continue

            chunks = list(paragraph_run_chunks(p))
            text_parts: list[str] = []

            def flush():
                nonlocal text_parts
                if not text_parts:
                    return
                combined = sanitize_line("".join(text_parts))
                text_parts = []
                if not combined:
                    return
                if is_junk_numbered_bullet(combined):
                    return
                is_list = p.style and "List" in (p.style.name or "")
                lines.append("")
                if is_list:
                    item = combined.lstrip("•").strip()
                    if item:
                        lines.append(f"- {item}")
                else:
                    lines.append(combined)

            for kind, payload in chunks:
                if kind == "image":
                    flush()
                    url = copy_image(payload)
                    if url:
                        lines.append("")
                        lines.append(f"![]({url})")
                        lines.append("")
                else:
                    text_parts.append(payload)

            flush()
            continue

        if isinstance(block, Table):
            lines.append("")
            for row in block.rows:
                cells = []
                for c in row.cells:
                    ct = " ".join(
                        sanitize_line(paragraph_plain_text(pp)) for pp in c.paragraphs
                    ).strip()
                    cells.append(ct.replace("|", "\\|"))
                if any(cells):
                    lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

    # Collapse 3+ blank lines to 2
    cleaned: list[str] = []
    blank_run = 0
    for line in lines:
        if line.strip() == "":
            blank_run += 1
            if blank_run <= 2:
                cleaned.append(line)
        else:
            blank_run = 0
            cleaned.append(line)

    body = "\n".join(cleaned).strip() + "\n"
    title = title_override or title_guess or docx_path.stem
    if body and not body.lstrip().startswith("#"):
        body = f"# {title}\n\n" + body
    else:
        # Ensure single H1 matches sidebar title when overridden
        if title_override and body.lstrip().startswith("#"):
            bl = body.split("\n")
            if bl[0].startswith("#"):
                bl[0] = f"# {title_override}"
                body = "\n".join(bl)
    # Drop repeated title line immediately under H1 (Word often duplicates title as body text)
    bl = body.split("\n")
    if len(bl) >= 3 and bl[0].startswith("# ") and bl[1].strip() == "":
        h1 = bl[0].lstrip("#").strip()
        if bl[2].strip() == h1:
            bl = bl[:2] + bl[3:]
            body = "\n".join(bl)

    fm = f"""---
title: {title}
sidebar_label: {title}
---

"""
    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(fm + body, encoding="utf-8")
    print("Wrote", out_md)


def pdf_to_markdown_pages(
    pdf_path: Path,
    out_md: Path,
    static_img_subdir: str,
    doc_key: str,
    repo_root: Path,
    title_override: str | None = None,
    dpi: float = 144.0,
) -> None:
    import fitz

    dest_dir = repo_root / "docusaurus/static/img" / static_img_subdir
    dest_dir.mkdir(parents=True, exist_ok=True)

    doc = fitz.open(pdf_path)
    lines: list[str] = []
    title = title_override or pdf_path.stem.replace("-", " ").strip()

    for i in range(doc.page_count):
        page = doc[i]
        mat = fitz.Matrix(dpi / 72, dpi / 72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        fname = f"{doc_key}-page-{i+1:02d}.png"
        out_png = dest_dir / fname
        pix.save(str(out_png))
        url = f"/img/{static_img_subdir}/{fname}".replace("\\", "/")
        if i == 0:
            lines.append(f"![]({url})")
        else:
            lines.append("")
            lines.append(f"![]({url})")

    body = "\n".join(lines) + "\n"
    fm = f"""---
title: {title}
sidebar_label: {title}
---

# {title}

"""
    out_md.parent.mkdir(parents=True, exist_ok=True)
    out_md.write_text(fm + body, encoding="utf-8")
    print("Wrote", out_md, "pages", doc.page_count)


def main():
    repo_root = Path(__file__).resolve().parents[1]
    import os

    os.chdir(repo_root)

    base = Path(r"C:\Users\dbalan\OneDrive\Desktop\GURU CARD CONTENT\SOP's\Social&Content")

    docx_to_markdown(
        base / "Standard Process Outline.docx",
        repo_root / "docusaurus/docs/sop/social-content/standard-process-outline/index.md",
        "sop-social-content/standard-process-outline",
        "spo",
        repo_root,
    )
    docx_to_markdown(
        base / "Email Marketing Process Mapping.docx",
        repo_root / "docusaurus/docs/sop/social-content/email-marketing-process-mapping/index.md",
        "sop-social-content/email-marketing-process-mapping",
        "em",
        repo_root,
        title_override="Email Marketing Process Mapping",
    )
    pdf_to_markdown_pages(
        base / "SMM Standard and Plus - Service Expectations.pdf",
        repo_root / "docusaurus/docs/sop/social-content/smm-standard-and-plus-service-expectations/index.md",
        "sop-social-content/smm-standard-and-plus-service-expectations",
        "smm",
        repo_root,
        title_override="SMM Standard and Plus - Service Expectations",
    )


if __name__ == "__main__":
    main()
